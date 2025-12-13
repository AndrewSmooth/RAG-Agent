import os
import json
from typing import List, Dict, Any
from pathlib import Path

from langchain_core.documents import Document
import hashlib
import fitz
from docx import Document as DocxDocument

from src.utils.clients import get_chroma_client
from src.constants import DOCS_COLLECTION_NAME, SQL_EXAMPLES_COLLECTION_NAME, T2T_DOCS_COLLECTION_NAME

class KnowledgeBaseLoader:
    def __init__(
        self,
        kb_path: str,
        chroma_url: str,
        yandex_api_key: str,
        yandex_folder_id: str,
    ):
        self.kb_path = kb_path
        self.chroma_url = chroma_url,
        self.yandex_api_key = yandex_api_key,
        self.yandex_folder_id = yandex_folder_id,
        self.chroma_client, self.embedding_fn = get_chroma_client(
            yandex_api_key=yandex_api_key,
            yandex_folder_id=yandex_folder_id
        )
        self._init_collections()

    def load_sql_examples(self, file: Any = None, filename: str = None) -> List[Document]:
        examples = []

        if file:
            document = self.get_document(file, filename, 'sql_example')
            examples.append(document)
        else:
            sql_dir = os.path.join(self.kb_path, SQL_EXAMPLES_COLLECTION_NAME)

            for filename in os.listdir(sql_dir):
                if filename.endswith(".json"):
                    filepath = os.path.join(sql_dir, filename)
                    with open(filepath, "r", encoding="utf-8") as f:
                        data = json.load(f)

                    examples.append(
                        Document(
                            page_content=f"Question: {data['question']}\nSQL: {data['sql']}",
                            metadata={
                                "source": filename,
                                "type": "sql_example",
                                "question": data["question"]
                            }
                        )
                    )
        return examples

    def load_docs(
            self,
            file: Any = None,
            filename: str = None,
            docs_type: str = T2T_DOCS_COLLECTION_NAME,
            docs_dir: str = T2T_DOCS_COLLECTION_NAME
    ) -> List[Document]:
        """Загружает документы в зависимости от параметров в docs или t2t_docs"""

        docs = []

        if file is not None:
            if not isinstance(file, bytes):
                raise TypeError("`file` must be bytes (result of UploadFile.read())")

            ext = Path(filename).suffix.lower()
            content_str = self._read_text_from_bytes(file, ext)

            if content_str.strip():
                document = Document(
                    page_content=content_str,
                    metadata={"source": filename, "type": docs_type}
                )
                docs.append(document)

        else:
            # Загрузка из директории
            docs_dir = os.path.join(self.kb_path, docs_dir)
            if not os.path.exists(docs_dir):
                return docs

            for fname in os.listdir(docs_dir):
                fpath = os.path.join(docs_dir, fname)
                if not os.path.isfile(fpath):
                    continue

                ext = Path(fname).suffix.lower()
                if ext not in {".txt", ".md", ".pdf", ".docx"}:
                    continue

                try:
                    with open(fpath, "rb") as f:  # читаем как байты для единообразия
                        raw = f.read()
                    content_str = self._read_text_from_bytes(raw, ext)

                    if content_str.strip():
                        document = Document(
                            page_content=content_str,
                            metadata={"source": fname, "type": docs_type}
                        )
                        docs.append(document)
                except Exception as e:
                    print(f"❌ Ошибка чтения {fname}: {e}")
                    continue

        return docs

    def load_file(
            self,
            file: Any,
            file_name: str,
            doc_type: str = None
    ):
        if not doc_type:
            return {'error': 'Не указан тип документа'}

        doc = []
        filename = file_name

        if doc_type == DOCS_COLLECTION_NAME:
            doc = self.load_docs(file, filename, docs_type='doc', docs_dir='docs')
        elif doc_type == SQL_EXAMPLES_COLLECTION_NAME:
            doc = self.load_sql_examples(file, filename)
        elif doc_type == T2T_DOCS_COLLECTION_NAME:
            doc = self.load_docs(file, filename)
        else:
            return {'error': f'Указан недопустимый тип документа: {doc_type}'}

        doc_id = f'{doc_type}_{self.hash_filename(filename)}'
        doc_text = doc[0].page_content

        chroma_collection = self.chroma_client.get_collection(doc_type)
        existing = chroma_collection.get(ids=[doc_id])

        embedding = self.embedding_fn(doc_text).tolist()

        if existing['ids']:
            print(f"⚠️ Документ с ID '{doc_id}' уже существует. Пропускаем.")
        else:
            chroma_collection.add(
                ids=[doc_id],
                documents=[doc_text],
                metadatas=[doc[0].metadata],
                embeddings=[embedding], # явно добавляем. чтобы не было багов
            )
            print(f"✅ Добавлен новый документ: {doc_id}")

        return {'ok': True}

    @staticmethod
    def hash_filename(filename: str, length: int = 16) -> str:
        return hashlib.sha256(filename.encode()).hexdigest()[:length]

    @staticmethod
    def get_document(file_content: Any, file_name: str, doc_type: str,) -> Document:
        document = Document(
            page_content=file_content,
            metadata={
                "source": file_name,
                "type": doc_type,
            }
        )
        return document

    def _init_collections(self):
        collections_names = [
            DOCS_COLLECTION_NAME,
            SQL_EXAMPLES_COLLECTION_NAME,
            T2T_DOCS_COLLECTION_NAME
        ]

        for collection_name in collections_names:
            self.chroma_client.get_or_create_collection(name=collection_name)

    @staticmethod
    def _read_text_from_bytes(content: bytes, ext: str) -> str:
        """Читает текст из байтов в зависимости от расширения."""
        try:
            if ext == ".pdf":
                with fitz.open(stream=content, filetype="pdf") as doc:
                    return "".join(page.get_text() for page in doc)
            elif ext == ".docx":
                from io import BytesIO
                doc = DocxDocument(BytesIO(content))
                return "\n".join(para.text for para in doc.paragraphs)
            else:
                # .txt, .md и прочие текстовые — декодируем как UTF-8
                return content.decode("utf-8")
        except UnicodeDecodeError:
            # Попытка с игнорированием ошибок для текстовых файлов
            return content.decode("utf-8", errors="replace")
